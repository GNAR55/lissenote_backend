from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
import itertools
import re
import os.path
from typing import List


def write_paragraph(doc: Document, content: str) -> Paragraph:
    """
    Writes a new paragraph to document.
    
    :param doc: Document
    :param content: Content to be written
    :return: None
    """
    return doc.add_paragraph(content)


def write_text(para: Paragraph, content: str):
    """
    Writes to document.

    :param para: Paragraph
    :param content: Content to be written
    :return: None
    """
    para.add_run(content)


def write_bold(para: Paragraph, content: str):
    """
    Writes in bold.

    :param para: Paragraph
    :param content: Content to be written
    :return: None
    """
    para.add_run(content).bold = True


def write_italic(para: Paragraph, content: str):
    """
    Writes in italic.
    
    :param para: Paragraph
    :param content: Content to be written
    :return: None
    """
    para.add_run(content).italic = True


def write_heading(doc: Document, heading: str, level: int):
    """
    Writes heading of specified level into document.
    
    :param doc: Document
    :param heading: Content of heading
    :param level: Level of heading
    :return: None
    """
    doc.add_heading(heading, level=level)


def write_caption(doc: Document, caption: str):
    """
    Write caption of image into document.
    
    :param doc: Document
    :param caption: Caption of image
    :return: None
    """
    p = doc.add_paragraph()
    write_italic(p, "figure: " + caption)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def write_image(doc: Document, image: str, caption: str, width: float = 2):
    """
    Writes image to document along with caption.
    
    :param doc: Document
    :param image: Path to image
    :param caption: Caption of image
    :param width: Width of image in the generated document
    :return: None
    """
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(image, width=Inches(width))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    write_caption(doc, caption)


def write_highlighted(doc: Document, content: str):
    """
    Writes highlighted text.
    
    :param doc: Document
    :param content: Content to be highlighted
    :return: None
    """
    doc.add_run(content).font.highlight_color = WD_COLOR_INDEX.YELLOW


def write_hyperlink(doc: Document, url: str):
    """
    Writes hyperlink into document.
    
    :param doc: Document
    :param url: Complete URL
    :return: None
    """
    run = doc.add_paragraph().add_run(url)
    run.underline = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xEE)


def write_keyword(para: Paragraph, keyword: str):
    """
    Write keyword into document.
    
    :param para: Paragraph
    :param keyword: Important term
    :return: None
    """
    run = para.add_run(keyword + " ")
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def separate_keywords(text: str, keywords: List[str]) -> List[str]:
    """
    Separates keywords from text.

    :param text: Paragraph content
    :param keywords: Important term
    :return: List with text separated form keywords
    """
    indices = []
    for keyword in keywords:
        indices += [i.start() for i in re.finditer(keyword, text)]
        indices += [i.end() for i in re.finditer(keyword, text)]

    indices = sorted(indices) if 0 in indices else [0] + sorted(indices)
    offsetted_indices = indices[1:]+[None]
    return [text[i:j] for i, j in zip(indices, offsetted_indices)]


def write_section(doc: Document, text: str, keywords: List[str], image_content: List[str] = []):
    """
    Writes one section consisting of text and a related image.

    :param doc: Document
    :param text: Text in a paragraph
    :param keywords: Important terms
    :param image_content: List containing image path followed by caption
    :return: None
    """
    p = write_paragraph(doc, "")
    if text:
        for item in separate_keywords(text, keywords):
            if item in keywords:
                write_keyword(p, item)
            else:
                write_text(p, item)

    if image_content:
        image, caption = image_content
        write_image(doc, image, caption)
    return p


def to_docx(topic: str, paragraphs: List[str], keywords: List[str], image_content: List[List[str]], links: List[str], output_directory='.', output_filename: str = 'converted'):
    """
    Generates a docx file.

    :param str topic: Topic of the notes
    :param paragraphs: Paragraphs
    :param keywords: Important terms
    :param image_content: Nested list with each element containing path to image followed by caption
    :param links: Links to websites/webpages related to the topic
    :param output_directory: Directory in which generated notes are to be saved in
    :param output_filename: Path of output docx file
    :returns: Notes in docx file
    :rtype: Document
    """
    docx = Document()
    write_heading(docx, topic, 0)

    for paragraph, image_content in list(itertools.zip_longest(paragraphs, image_content)):
        write_section(docx, paragraph, keywords, image_content)

    write_heading(docx, "Related Links", 1)
    for link in links:
        write_hyperlink(docx, link)

    docx.save(os.path.join(output_directory, output_filename + '.docx'))
    return docx
